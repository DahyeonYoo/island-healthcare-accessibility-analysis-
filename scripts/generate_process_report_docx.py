# -*- coding: utf-8 -*-
"""단계별 과정 설명 보고서 + 학습 액션플랜 docx 생성"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _set_run_font(run, east_asia="휴먼명조", size=13, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), east_asia)

from project_paths import BASE
OUT = BASE / "03_master" / "보고서_과정설명_및_학습액션플랜.docx"


def set_doc_style(doc):
    sec = doc.sections[0]
    sec.top_margin = Mm(15)
    sec.bottom_margin = Mm(15)
    sec.left_margin = Mm(20)
    sec.right_margin = Mm(20)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "휴먼명조")
    style.font.size = Pt(13)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.6


def add_title(doc, text, size=20):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, east_asia="헤드라인M", size=size, bold=True)
    p.paragraph_format.space_after = Pt(12)


def add_heading(doc, text, level=1):
    sizes = {1: 15, 2: 13, 3: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, east_asia="헤드라인M" if level == 1 else "휴먼명조",
                  size=sizes.get(level, 13), bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def add_sub(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("□ " + text)
    _set_run_font(run, bold=True)


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run)
    p.paragraph_format.first_line_indent = Mm(0)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_run_font(run)
    p.paragraph_format.left_indent = Mm(5 + level * 5)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(40, 40, 40)
    p.paragraph_format.left_indent = Mm(5)


def build():
    doc = Document()
    set_doc_style(doc)

    add_title(doc, "전남 유인도 의료 접근성 연구")
    add_body(doc, "단계별 전처리·통계분석 과정 설명 보고서 및 학습 액션플랜")
    add_body(doc, "작성일: 2026-06-23  |  범위: 데이터 통합 ~ STEP3 공간 시각화")
    doc.add_page_break()

    # ── PART 0: 프로젝트 진행 종합 ──
    add_heading(doc, "0. 프로젝트 진행 내용 종합", 1)
    add_sub(doc, "연구 목표")
    add_body(doc,
        "전라남도 277개 유인도를 대상으로, 응급의료·1차의료에 대한 '총 도달시간(Y)'을 "
        "277섬 전원에 대해 산출하고, 그 원인·유형·공간 패턴을 분석하여 정책 우선순위를 "
        "제시하는 것이 본 프로젝트의 목표이다. 분석은 3단계로 설계되었다.")
    add_bullet(doc, "STEP 0~E: 데이터 전처리 및 복합 종속변수 Y 구축 (277섬 × 181열)")
    add_bullet(doc, "STEP 1: 원인 분석 — OLS + Random Forest + SHAP + 병목 분해")
    add_bullet(doc, "STEP 2: K-means 군집 — 취약도 3계층(C0/C1/C2) 분류")
    add_bullet(doc, "STEP 3: 공간 시각화 — 군집·병목유형·Y강도·이중취약 지도")
    add_bullet(doc, "(예정) 센터 데이터: 격자인구·SKT·카드 → 수요가중 취약지수 보강")

    add_sub(doc, "핵심 설계 결정")
    add_bullet(doc, "분석 기준키: 읍면동이 아닌 500m 격자(gid)로 통일, 최종 보고 단위는 섬(섬코드)")
    add_bullet(doc, "좌표계: EPSG:5179 (국토통계 접근성 shp 기준)")
    add_bullet(doc, "Y 정의: min(섬내 도달시간, sea_leg + land_leg) — 277섬 결측 0")
    add_bullet(doc, "sea_leg = 해상거리 ÷ 30km/h × 60 (ferry 실측은 노이즈로 Y에 미사용)")
    add_bullet(doc, "land_leg = 최근접 육지격자의 도로 접근성(km) ÷ 40km/h × 60 (분)")
    add_bullet(doc, "X에서 Y 구성요소(거리·페리·섬내의료·접근성) 제외 → 누수 방지")

    add_sub(doc, "주요 산출물 (03_master/)")
    add_bullet(doc, "island_analysis.csv / master_unified.csv — 최종 분석 테이블")
    add_bullet(doc, "step1_X.csv, step1_model_report_v2.txt — STEP1 모델 결과")
    add_bullet(doc, "stepA_bottleneck.csv — 병목 유형(육지/해상/혼합/자립)")
    add_bullet(doc, "step2 클러스터링/step2_clusters.csv — K-means 3군집")
    add_bullet(doc, "step3 공간 시각화/*.png — 4종 지도")

    add_sub(doc, "해결한 주요 기술 이슈")
    add_bullet(doc, "섬 경계 shp에 행정 섬코드 없음 → 좌표 공간조인(within→nearest)으로 crosswalk")
    add_bullet(doc, "shp_id 인덱스 불일치 버그 → 전남 필터 후 reset_index로 통일")
    add_bullet(doc, "국토통계 접근성 -999(도로 미연결) → 216/277섬 무자료 → Y 재정의로 전수 커버")
    add_bullet(doc, "ferry 실측 노이즈(r=0.20) → 거리기반 sea_leg로 일관 산출")
    add_bullet(doc, "접근성 단위(km vs 분) 혼동 → step_f에서 km→분 환산(도로 40km/h) 확정")

    doc.add_page_break()

    # ── PART 1: 보고서 양식 개요 ──
    add_heading(doc, "1. 대회 보고서 양식 분석 및 우리 연구 매핑 개요", 1)
    add_body(doc,
        "『2026년 국가데이터 활용대회 안내문』 8~12쪽(붙임4)에 따르면, "
        "데이터분석 보고서는 7쪽 내외, 한글(HWP) 제출, "
        "SDC/MDIS 제공자료 1종 이상 필수 활용이 요구된다. "
        "본 문서는 최종 제출용 7쪽 보고서의 목차·내용 배치 가이드이다.")

    add_heading(doc, "1.1 공식 보고서 구조 (안내문 p.9)", 2)
    add_bullet(doc, "제목 [헤드라인M, 20pt]")
    add_bullet(doc, "1. 배경 — □ 주제 선정 / □ 분석 필요성(문제점) 및 전략")
    add_bullet(doc, "2. 데이터 분석 — □ 데이터 선정 / □ 데이터 분석(프로세스·방법) / □ 분석 결과 및 해석")
    add_bullet(doc, "3. 분석 활용 전략 — □ 기대효과 / □ 방향제시")
    add_bullet(doc, "참고문헌 (p.10, 10pt)")

    add_heading(doc, "1.2 우리 연구 → 보고서 7쪽 배치안", 2)

    sections = [
        ("제목 (약 0.5쪽)",
         "예: 「전남 유인도 의료 접근성 취약지 진단과 수요가중 정책 우선순위」\n"
         "핵심 메시지 한 줄: '277개 유인도 전수 분석으로 해상·육지 병목을 분해하고, "
         "고령 수요를 반영한 개입 우선순위를 제시'"),
        ("1. 배경 (약 1.5쪽)",
         "□ 주제 선정\n"
         "  - 전남 277유인도, 국토통계 접근성은 배로만 가는 섬 81% 무자료\n"
         "  - 기존 지표는 '본토에서 멀다'만 보여 수요(고령·실거주) 반영 부족\n"
         "  - SDC 격자인구·소지역·카드/SKT 연계로 공급-수요 통합 필요\n"
         "□ 분석 필요성 및 전략\n"
         "  - 문제: 응급 도달시간 산출 불가 섬 216개 → 정책 대상 누락\n"
         "  - 전략: 격자→섬 공간통합 → 복합 Y → 병목분해 + 군집 + 수요가중 지수"),
        ("2. 데이터 분석 — 데이터 선정 (약 0.8쪽)",
         "표 1 권장: 데이터명 | 출처 | 시계열 | 역할\n"
         "  - SDC 격자단위 인구(2020~2024) [필수] — 고령·학령·청년 수요\n"
         "  - SDC 소지역통계 / SKT 실거주 / 카드매출 [필수·선택] — 섬 단위 집계\n"
         "  - 국토통계 접근성 shp 12종(2020~2023) — land_leg\n"
         "  - 전라남도 유인도정보 CSV — 섬 속성·인프라 X\n"
         "  - AL_D158 섬경계 shp — 격자↔섬 매핑\n"
         "  - 여객선 운항(보조·검증용)"),
        ("2. 데이터 분석 — 분석 프로세스 (약 1.2쪽)",
         "그림 1 권장: 파이프라인 흐름도 (Step B→C→D→E→1→2→3)\n"
         "  - 전처리: crosswalk, 격자병합, Y=sea+land, 단위정합(km→분)\n"
         "  - 통계: OLS(VIF), RF+SHAP, 병목분해, K-means(k=3), 공간지도\n"
         "  - 센터: prep01~05 → island_analysis_plus → 수요가중 지수"),
        ("2. 데이터 분석 — 결과 및 해석 (약 2.0쪽)",
         "핵심 수치 3~4개 + 지도 2~3개 (7쪽 분량 고려)\n"
         "  - Y: 평균 62.3분, 가거도 236분 / 진도 3.4분\n"
         "  - 병목: 육지지배 198개(71%), 해상지배 27개(10%)\n"
         "  - SHAP(먼섬특별법 제외): 소규모급수·인구밀도·해수담수화 = 고립 proxy\n"
         "  - K-means: C0(146) 양호 / C1(117) 이중취약 / C2(14) 극취약\n"
         "  - (센터 후) 응급_고령부담 Top15 — 정책 1순위 섬"),
        ("3. 분석 활용 전략 (약 1.0쪽)",
         "□ 기대효과: 섬별 맞춤 개입(항로·권역응급·순회진료) / 예산 효율\n"
         "□ 방향제시:\n"
         "  - C2(외해): 응급헬기·병원선·항로 증편\n"
         "  - C1(중취약·인구과소): 순회·이동형 의료·응급이송\n"
         "  - C0(본토근접): 육지 거점 의료망 연계\n"
         "  - 한계: ferry 30km/h·도로 40km/h 가정 → 민감도 분석 권장"),
    ]
    for title, content in sections:
        add_sub(doc, title)
        add_body(doc, content)

    add_heading(doc, "1.3 심사 기준 대응 포인트", 2)
    add_bullet(doc, "주제 독창성: 277섬 전수 + sea/land 병목 분해 (기존 접근성 지표 한계 극복)")
    add_bullet(doc, "데이터 활용성: SDC 격자인구·소지역 필수 + 공공데이터 다층 연계")
    add_bullet(doc, "분석 차별성: Y 재정의, 누수 제어, 예측(SHAP) vs 개입(병목) 이중 프레임")
    add_bullet(doc, "결과 활용성: C0/C1/C2 × 병목유형 → 구체적 정책 레버 매핑")

    doc.add_page_break()

    # ── PART 2: 단계별 과정 설명 ──
    add_heading(doc, "2. 단계별 전처리·통계분석 과정 설명", 1)

    steps = [
        ("Step B-1: 섬코드 ↔ 섬경계 Crosswalk (step_b_crosswalk.py)",
         "목적: AL_D158 섬경계 shp에는 행정 '섬코드'가 없어 CSV와 직접 join 불가.\n"
         "논리:\n"
         "  1) 유인도 CSV의 위·경도를 EPSG:5179 점으로 변환\n"
         "  2) 1순위 point-in-polygon(within): 점이 속한 전남 섬 polygon 선택\n"
         "  3) 2순위 nearest(≤5km): within 실패 시 최근접 polygon\n"
         "  4) 신뢰도: within=high, nearest<500m=medium, ≥500m=low, 미매칭=none\n"
         "결과: 275/277 매칭 (high 241, medium 19, low 15, none 2)\n"
         "산출: 02_interim/island_crosswalk.csv\n"
         "핵심 코드: gpd.sjoin(pts, jn_keep, predicate='within')"),
        ("Step B-2: 섬 경계 구축 (step_b_boundaries.py)",
         "목적: 277섬 전체에 usable polygon/buffer geometry 확보.\n"
         "논리:\n"
         "  - high/medium crosswalk → shp polygon 사용\n"
         "  - low/none → CSV 좌표+면적 기반 원형 버퍼\n"
         "  - 면적 비율(shp/csv)이 0.3~3배 밖이면 버퍼로 강등 (접도→진도 흡수 오류 차단)\n"
         "산출: 02_interim/island_boundaries_jeonnam.gpkg (shp 251 + buffer 26)"),
        ("Step B-3: 격자↔섬 매핑 (step_b_grid_lookup.py)",
         "목적: 49,233 전남 격자 중 각 섬에 귀속되는 격자 식별.\n"
         "논리:\n"
         "  - 격자 centroid가 섬 polygon 내부 → 해당 섬 귀속\n"
         "  - 섬보다 작은 섬: 대표점이 속한 격자 1칸 강제 배정\n"
         "  - long-format(다대다): 25개 공유격자 허용\n"
         "결과: 277/277 섬 ≥1 격자, 귀속 격자 6,755개\n"
         "산출: 02_interim/grid_island_lookup.csv"),
        ("Step C: 접근성 12레이어 병합 (step_c_merge_access.py)",
         "목적: 보건기관/응급의료/의원 × 2020~2023 value를 gid 기준 wide merge.\n"
         "논리: value==-999 → NaN. 전남 clip(sido_cd=46).\n"
         "산출: grid_access_table.csv, grid_master_jeonnam.gpkg"),
        ("Step D: 섬 단위 분석 테이블 (step_d_island_table.py)",
         "목적: STEP1 입력용 island_analysis.csv 생성.\n"
         "논리:\n"
         "  1) grid_access × grid_island_lookup → 섬별 mean/min/max 집계\n"
         "  2) + 유인도 CSV 전체 속성\n"
         "  3) + ferry_island_access (검증·보조)\n"
         "산출: 03_master/island_analysis.csv (277 × 170+열)"),
        ("Step E: 복합 종속변수 Y (step_e_final_Y.py, step_f_fix_and_merge.py)",
         "목적: 277섬 전원에 대해 Y_time_emergency/clinic 산출 (결측 0).\n"
         "Y 정의:\n"
         "  sea_leg(분) = 0 (육지연결) | dist_main_km / 30 × 60 (미연결)\n"
         "  land_leg(분) = 최근접 육지(비섬)격자 acc_*_2023(km) / 40 × 60\n"
         "  Y = min(섬내_도달시간(시설 보유), sea_leg + land_leg)\n"
         "단위 정합(step_f): 국토통계 value=도로거리(km) 확인 → km를 분으로 환산.\n"
         "팀원 X변수 병합 → master_unified.csv\n"
         "분포: Y_emergency 평균 62.3분, 중앙값 54.3, max 236.4(가거도)"),
        ("STEP 1-A: EDA & 전처리 (step1_eda.py)",
         "목적: 모델 입력 step1_X.csv 생성.\n"
         "논리:\n"
         "  - Y 왜도 1.51 → log1p 변환 (왜도 -0.75)\n"
         "  - X 76개: 인구·면적·지역지정(_b)·생활인프라·교육\n"
         "  - 누수 제외: Y 구성요소(거리·페리·섬내의료·acc_·연결성) 패턴 매칭\n"
         "  - 결측: 증감률_num 4건 → 중앙값 대치\n"
         "산출: step1_X.csv, step1_eda_report.txt"),
        ("STEP 1-B: OLS + RF + SHAP (step1_model.py, step1_model_v2.py)",
         "목적: Y에 영향을 주는 변수 탐색.\n"
         "(A) OLS: StandardScaler → VIF>10 반복 제거 → adjR²=0.455~0.478\n"
         "  유의(악화): 소규모급수·디젤발전·개발대상섬\n"
         "  유의(개선): NH농협·119·인구밀도·풍력발전\n"
         "(B) RF: OOB R²=0.463~0.626 (비선형 존재)\n"
         "(C) SHAP: 평균|SHAP| 상위 — 먼섬특별법(동어반복) 제외 후\n"
         "  실질 Top3: 소규모급수시설_b, 인구밀도, 해수담수화시설_b\n"
         "  → '고립도 proxy'로 재해석 (의료 직접 원인 아님)"),
        ("STEP 1-C: 병목 분해 (step2a_bottleneck.py)",
         "목적: '무엇을 고쳐야 하는가' — Y를 sea_leg + land_leg로 분해.\n"
         "분류 규칙:\n"
         "  - 섬내 시설이 Y 결정 → 섬내자립형 (2개)\n"
         "  - sea_share ≥ 60% → 해상지배형 (27개)\n"
         "  - sea_share ≤ 40% → 육지지배형 (198개)\n"
         "  - 그 외 → 혼합형 (50개)\n"
         "핵심 발견: 평균 land_leg(41.6분) > sea_leg(20.9분) → 육지측 병목 우세\n"
         "  그러나 Top 취약(가거도·홍도)은 전부 해상지배형\n"
         "산출: stepA_bottleneck.csv"),
        ("STEP 2: K-means 군집 (step2b_kmeans.py)",
         "목적: 정책 3계층(양호/중취약/극취약) 분류.\n"
         "피처(5): sea_leg, land_leg, Y, 인구밀도, dist_main_km → StandardScaler\n"
         "k 선택: silhouette 최대 k=2, 정책 해석 위해 k=3 채택\n"
         "  C0(n=146): Y=36분, 본토근접\n"
         "  C1(n=117): Y=81분, land_leg↑ + 인구밀도↓ → 이중취약\n"
         "  C2(n=14): Y=177분, 외해 고립(가거도·홍도·흑산)\n"
         "교차검증: C2는 해상지배형 11/14\n"
         "산출: step2_clusters.csv"),
        ("STEP 3: 공간 시각화 (step3_maps.py)",
         "목적: 분석 결과의 공간 패턴·정책 메시지 시각화.\n"
         "  map1: 군집 3계층 (동→서 취약도 증가)\n"
         "  map2: 병목 유형 (육지지배 광범위, 해상지배 서쪽 외해)\n"
         "  map3: Y 강도 단계구분\n"
         "  map4: C1 이중취약 — 마커 크기=정주인구\n"
         "좌표: EPSG:5179, island_boundaries_jeonnam.gpkg 배경"),
        ("(예정) 센터 데이터 통합 (prep01~05)",
         "prep01_discover: STAT_ITM 코드·GRID↔gid 매칭·좌표계·결합키 검증\n"
         "prep02: 소지역 좌표 → 섬 polygon 공간조인 → BLOCK↔섬코드\n"
         "prep03: SDC 격자인구 → grid_island_lookup → 섬 집계\n"
         "prep04: 카드·SKT → 소지역 룩업 → 섬 집계\n"
         "prep05: island_analysis_plus + 응급_고령부담 = 고령인구 × Y"),
    ]

    for i, (title, content) in enumerate(steps, 1):
        add_heading(doc, f"2.{i} {title}", 2)
        for para in content.split("\n"):
            if para.strip().startswith("-"):
                add_bullet(doc, para.strip().lstrip("- "))
            elif para.strip().startswith("  "):
                add_bullet(doc, para.strip(), level=1)
            else:
                add_body(doc, para)

    doc.add_page_break()

    # ── PART 3: 학습 액션플랜 ──
    add_heading(doc, "3. 학습 액션플랜 — 개념·논문·코드 체득 로드맵", 1)

    add_heading(doc, "3.1 통계·분석 기법 학습 순서 (4주 권장)", 2)
    weeks = [
        ("1주차: 기초 통계 + 탐색",
         ["기술통계, 상관분석, 왜도·로그변환(log1p) — step1_eda.py",
          "결측 처리(중앙값 대치), 이진화(_b), 파생변수(인구_log, 인구밀도)",
          "실습: step1_X.csv로 Y 분포·상관표 재현"]),
        ("2주차: 회귀 + 공선성",
         ["OLS, 표준화, VIF, adjR² 해석 — step1_model.py 37~65행",
          "누수(leakage) 개념: Y 구성요소를 X에 넣으면 안 되는 이유",
          "실습: VIF 반복 제거 전후 계수 변화 비교"]),
        ("3주차: 머신러ning + SHAP",
         ["Random Forest, OOB score, 5-fold CV — step1_model.py 67~76행",
          "SHAP TreeExplainer, mean|SHAP| 해석 — step1_model_v2.py",
          "동어반복(tautology): 먼섬특별법 사례로 '원인 vs 라벨' 구분",
          "실습: 먼섬특별법 포함/제외 SHAP 비교"]),
        ("4주차: 군집 + 공간분석",
         ["StandardScaler, K-means, silhouette — step2b_kmeans.py",
          "GeoPandas: sjoin, sjoin_nearest, CRS 변환 — step_b_crosswalk.py",
          "Choropleth 지도 — step3_maps.py",
          "실습: k=2~6 silhouette 그래프 직접 그리기"]),
    ]
    for wtitle, items in weeks:
        add_sub(doc, wtitle)
        for item in items:
            add_bullet(doc, item)

    add_heading(doc, "3.2 신뢰할 수 있는 참고문헌", 2)
    refs = [
        "Guagliardo, M. F. (2004). Spatial accessibility of primary care: Concepts, methods and challenges. Primary Health Care Research & Development, 5(3), 197-205. — 접근성 개념의 고전",
        "Luo, W., & Wang, F. (2003). Measures of spatial accessibility to health care in a GIS environment. International Journal of Health Geographics, 2, 3. — GIS 기반 의료 접근성",
        "Lundberg, S. M., & Lee, S.-I. (2017). A Unified Approach to Interpreting Model Predictions (SHAP). NeurIPS. — SHAP 원논문",
        "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5-32. — RF 원논문",
        "국토교통부 (2022~2023). 국토정책지표 — 응급의료시설·의원·보건기관 접근성(500m 격자). 국토통계. — land_leg 데이터 정의",
        "통계청 (2024). SGIS 소지역통계·격자단위 인구통계 이용안내. — SDC 필수 데이터",
        "전라남도 (2024). 유인도 현황조사. — 섬 속성 원본",
        "James, G., et al. (2021). An Introduction to Statistical Learning (2nd ed.). Springer. — OLS/RF/K-means 입문 (무료 PDF)",
    ]
    for r in refs:
        add_bullet(doc, r)

    add_heading(doc, "3.3 코드별 집중 학습 포인트", 2)
    code_focus = [
        ("step_b_crosswalk.py", "공간조인(within/nearest) — GIS+통계 연결의 핵심. 57~60행 sjoin 로직"),
        ("step_e_final_Y.py", "Y 정의 전체. 40~80행 sea_leg/land_leg/Y 산출 — 연구의 심장"),
        ("step_f_fix_and_merge.py", "km→분 단위 정합. 30~48행 — '단위 혼동' 디버깅 사례"),
        ("step1_eda.py", "누수 변수 제외 패턴(37~47행). 통계 모델링 전 필수 전처리"),
        ("step1_model_v2.py", "VIF→OLS→RF→SHAP 파이프라인 한 파일에 집약"),
        ("step2a_bottleneck.py", "32~44행 classify() — 통계+정책 연결의 모범"),
        ("step2b_kmeans.py", "41~61행 k 탐색·재라벨링 — '통계 vs 해석' trade-off"),
        ("step3_maps.py", "47~80행 지도 4종 — 보고서 Figure 제작"),
        ("prep01_discover.py", "센터 데이터 4가지 검증 — 새 데이터 반입 시 첫 스크립트"),
    ]
    for fname, desc in code_focus:
        add_bullet(doc, f"{fname}: {desc}")

    add_heading(doc, "3.4 핵심 개념 체크리스트 (스스로 설명할 수 있어야 함)", 2)
    checklist = [
        "왜 읍면동이 아니라 500m 격자를 썼는가?",
        "Y = min(섬내, sea+land)에서 ferry 실측 대신 거리기반 sea_leg를 쓴 이유는?",
        "국토통계 -999가 216섬에 나타나는 구조적 이유는?",
        "SHAP 상위 변수가 '의료 원인'이 아니라 '고립 proxy'인 이유는?",
        "병목 분해(육지/해상)와 SHAP(급수·인구밀도)의 역할 차이는?",
        "C1 이중취약과 C2 극취약의 정책 레버 차이는?",
        "응급_고령부담 = 고령인구 × Y가 '왜 이 섬부터?'에 답하는 방식은?",
    ]
    for c in checklist:
        add_bullet(doc, c)

    add_heading(doc, "3.5 실습 과제 (체득용)", 2)
    tasks = [
        "island_analysis.csv에서 Y_time_emergency 상위 10섬을 sea/land로 분해해 표 작성",
        "step1_model_v2.py 실행 후 SHAP summary plot 해석 1페이지 작성",
        "ROAD_SPEED를 30/40/50으로 바꿔 Y 분포 변화 → 민감도 1문단",
        "step3 지도 4종 중 1장을 PPT용으로 재편집 + 캡션 3줄",
        "(센터 후) prep05 실행 → 응급_고령부담 Top15와 C1 교차표 작성",
    ]
    for t in tasks:
        add_bullet(doc, t)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
